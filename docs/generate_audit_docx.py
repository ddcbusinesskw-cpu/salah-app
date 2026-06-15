#!/usr/bin/env python3
"""Generate audit-report.docx from the comprehensive audit findings."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── RTL support ─────────────────────────────────────────────────────────────
def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def rtl_para(text, style='Normal', bold=False, size=None, color=None):
    p = doc.add_paragraph(style=style)
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    set_run_rtl(run)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def heading(text, level=1):
    h = doc.add_heading(level=level)
    set_rtl(h)
    run = h.add_run(text)
    set_run_rtl(run)
    if level == 1:
        run.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x18, 0x4e, 0x4a)
    else:
        run.font.color.rgb = RGBColor(0x37, 0x4a, 0x47)
    return h

def table_row(table, cells, bold_first=False, header=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = ''
        p = cell.paragraphs[0]
        set_rtl(p)
        run = p.add_run(text)
        set_run_rtl(run)
        run.bold = bold_first and i == 0 or header
        if header:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if header:
        for cell in row.cells:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '0F766E')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading)
    return row

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=0, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_row(t, headers, header=True)
    for row_data in rows:
        table_row(t, row_data)
    if col_widths:
        for i, row in enumerate(t.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])
    return t

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    set_rtl(p)
    pf = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:right'), str(400 + indent * 360))
    pf.append(ind)
    run = p.add_run(text)
    set_run_rtl(run)
    return p

# ── Page setup (RTL) ─────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Set document-level RTL
sectPr = section._sectPr
bidi = OxmlElement('w:bidi')
sectPr.append(bidi)

# Default font
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
heading('تقرير التدقيق الشامل', 1)
rtl_para('تطبيق «المجدّد — طريق الأتقياء»', bold=True, size=14)
rtl_para('DDC Business', size=11)
rtl_para('تاريخ التدقيق: 15 يونيو 2026', size=11)
rtl_para('الحالة النهائية: 🔴 غير جاهز للرفع — 12 مانعاً يجب إصلاحه', bold=True, size=12,
         color=(0xB9, 0x1C, 0x1C))
doc.add_paragraph()
rtl_para(
    'أُجري هذا التدقيق بواسطة فريق QA + أمان + خصوصية + امتثال + تصميم + '
    'مراجعة شرعية على الكود المصدري الكامل (index.html ≈ 6254 سطر + sw.js + '
    'ملفات الإعداد). النتائج مصنّفة إلى 🔴 مانع للنشر / 🟡 يُفضّل إصلاحه / 🟢 سليم.'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — BLOCKERS
# ════════════════════════════════════════════════════════════════════════════
heading('أولاً — الموانع 🔴 (12 مانعاً يجب إصلاحها قبل البناء)', 1)

blockers = [
    ('م١', 'sw.js السطر 23',
     'خطأ أسبقية عمليات في Service Worker: الشرط `req.headers.get(\'accept\')||\'\'` دائماً '
     'truthy → جميع الطلبات الفاشلة أوفلاين ترجع index.html بدل 503. '
     'ملف Firebase.js الفاشل يُرجع HTML فيتعطّل التطبيق كلياً.\n'
     'الحالي (خطأ):\n'
     'if(req.destination===\'document\'||req.headers.get(\'accept\')||\'\'includes(\'text/html\'))\n'
     'الصحيح:\n'
     'if(req.destination===\'document\'||(req.headers.get(\'accept\')||\'\').includes(\'text/html\'))'),
    ('م٢', 'index.html ~3669',
     'تسريب مستمع Auth: login(\'google\') تُسجّل onAuthStateChanged جديد عند كل استدعاء. '
     'إن أُغلق الـ popup فشلت المصادقة، المستمع لا يُلغى. كل محاولة تراكم مستمعاً → '
     'دخول مزدوج، وتطبيق نفس الدالة مرات متعددة لاحقاً.'),
    ('م٣', 'index.html ~5837',
     'حذف الحساب: cred.idToken غير موجود → يجب cred.credential.idToken. '
     'أي حذف يحتاج إعادة مصادقة على الجهاز الأصلي يفشل دائماً بـ [auth/invalid-credential].'),
    ('م٤', 'index.html ~2603',
     'showDua() مشروطة بإذن الإشعارات: '
     'الدعاء بعد الأذان لا يظهر إن كانت الإشعارات مُوقفة رغم عزف الأذان. '
     'showDua() يجب استدعاؤها خارج شرط Notification.permission.'),
    ('م٥', 'index.html ~2879',
     'buildQFlat() تجمّد الواجهة: معالجة 6236 آية بـ 7 regex لكل آية بشكل متزامن '
     'في main thread عند أول بحث. تجمّد 300–800ms على الأجهزة الضعيفة. '
     'الحل: requestIdleCallback أو Web Worker.'),
    ('م٦', 'index.html ~4236',
     'عداد الإفطار رمضان: المتغير TIMES غير معرَّف في أي مكان. '
     'الكود الصحيح: PT.data.timings.Maghrib. العداد صامت دائماً.'),
    ('م٧', 'index.html (كامل الملف)',
     'لا معالج لزر الرجوع في أندرويد: غياب App.addListener(\'backButton\'). '
     'الـ modals لا تُغلق، التطبيق يغادر مباشرة بلا تأكيد، '
     'المستخدم قد يعلق داخل نوافذ منبثقة.'),
    ('م٨', 'index.html ~5873',
     'XSS في _fbShowErr: e.message (رسالة Firebase متأثرة بمدخلات المستخدم) '
     'تُحقن في innerHTML مباشرةً. استبدل بـ textContent للجزء النصي.'),
    ('م٩', 'index.html ~3561',
     'Wake Lock لا يُحرَّر: مغادرة شاشة الركعات عبر go() لا تستدعي rakWakeLock(false). '
     'الشاشة تبقى مضاءة حتى تُغلق العملية.'),
    ('م١٠', 'index.html ~2885',
     'JSONP من dorar.net: يُحقن <script> من خادم خارجي في الـ DOM. '
     'إن اختُرق dorar.net: كود عشوائي يتنفّذ بكامل صلاحيات WebView '
     '(localStorage، Firebase tokens، حالة التطبيق). '
     'أيضاً: <script> وwindow[cb] لا يُنظَّفان عند الخطأ أو المهلة.'),
    ('م١١', 'build.gradle سطر 55 + AndroidManifest',
     'Facebook Login SDK محمَّل (v16.3.0) بـ placeholder credentials. '
     'Facebook SDK يجمع بيانات جهاز عند التهيئة. '
     'Google Play يرفض أو يُعلّق إن لم يُعلَن في Data Safety. '
     'الحل: احذف الـ dependency والـ meta-data إن لم يُستخدم.'),
    ('م١٢', 'index.html ~3683',
     'Apple Sign-In وهمي: login(\'apple\') تكتب session=\'apple\' وتستدعي enterApp() مباشرةً. '
     'لا Firebase auth، المستخدم يعتقد أن له حساباً لكنه ضيف. '
     'Apple تشترط Sign in with Apple حقيقياً إن وُجد تسجيل دخول اجتماعي آخر (Google). '
     'سبب رفض محتمل من App Store.'),
]

for code, loc, desc in blockers:
    heading(f'{code} — {loc}', 2)
    p = rtl_para(desc)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PREFERRED FIXES (table)
# ════════════════════════════════════════════════════════════════════════════
heading('ثانياً — الإصلاحات المُفضَّلة 🟡', 1)

preferred_headers = ['#', 'التصنيف', 'الموقع', 'الملاحظة']
preferred_rows = [
    ('ف١', 'أمان', 'index.html ~5603', 'SW registration .catch() فارغة — أضف console.warn كحد أدنى'),
    ('ف٢', 'أمان', 'index.html ~2767', 'Notification.requestPermission() بدون .catch() — بعض المتصفحات ترجع undefined'),
    ('ف٣', 'وظائف', 'index.html ~2856', 'clipboard.writeText() بدون .catch() — toast النجاح يظهر قبل استقرار Promise'),
    ('ف٤', 'وظائف', 'index.html ~4416,4445', 'navigator.share() بدون .catch() — إلغاء المستخدم يُطلق global error toast'),
    ('ف٥', 'وظائف', 'index.html ~2801', 'رفض إذن البوصلة على iOS: .catch() فارغة — الواجهة تبقى "حرّك الهاتف"'),
    ('ف٦', 'أمان', 'index.html ~2885', 'JSONP hadith: <script> وwindow[cb] لا يُنظَّفان عند timeout/error'),
    ('ف٧', 'بيانات', 'index.html ~3913', 'JSON.parse في Khatmah بدون try/catch — تلف localStorage يُعطّل الختمة'),
    ('ف٨', 'بيانات', 'index.html ~3820', 'Migration months→days لا تُحفظ بـ lsSet — تُعاد عند كل تشغيل'),
    ('ف٩', 'أداء', 'index.html ~4175', 'createObjectURL في مشاركة الإنجازات بلا revokeObjectURL — تسريب ذاكرة'),
    ('ف١٠', 'وظائف', 'index.html ~3683', 'Apple Sign-In (P-10): يجب تطبيق FirebaseAuthentication.signInWithApple() أو toast واضح'),
    ('ف١١', 'وظائف', 'index.html ~5536', 'iPadOS 13+ يُعيد "Macintosh" — regex الكشف لا يُغطّيه، زر Apple مخفي على iPad'),
    ('ف١٢', 'بيانات', 'index.html ~2820', 'openSurah() لا يتحقق من فشل lsSet بسبب امتلاء القرص'),
    ('ف١٣', 'وظائف', 'index.html ~2884', 'goAyah() 300ms delay يتسابق مع async fetch للسورة على اتصال بطيء'),
    ('ف١٤', 'أداء', 'index.html ~3499', 'التسبيح: كتابة localStorage في كل نقرة — عشرات الكتابات/ثانية. أضف debounce'),
    ('ف١٥', 'وظائف', 'index.html ~2612', 'نافذة الإشعار المبكر 5 دقائق ضيّقة جداً (5 ثوانٍ) — jitter قد يُخطئها'),
    ('ف١٦', 'وظائف', 'index.html ~2525', 'رفض إذن الموقع: انتقال صامت لـ city picker — أضف toast توضيحياً'),
    ('ف١٧', 'بيانات', 'index.html ~6178', 'مزامنة إعدادات Firebase: تُدمج بـ updatedAt فقط — تعارضات في نفس الثانية تُخسر تغييرات'),
    ('ف١٨', 'أداء', 'index.html ~6058', '_fbWriteAll يُشغَّل عند كل lsSet — حجم كتابة Firebase مرتفع جداً. قلّل إلى 30s'),
    ('ف١٩', 'أداء', 'index.html ~2353', 'المصحف الكامل: DOM يتراكم (114 سورة) بلا إزالة للبعيدة — ~5MB DOM'),
    ('ف٢٠', 'امتثال', 'AndroidManifest سطر 4', 'allowBackup="true" يسمح بنسخ localStorage لجهاز آخر'),
    ('ف٢١', 'هوية', 'DICT ~5545', 'مفتاح DICT الميت "نور بريميوم" → "Noor Premium" في الترجمات (غير مُستخدم حالياً)'),
    ('ف٢٢', 'توطين', 'index.html متعدد', '25 سلسلة نصية لفئات الأذكار/الأدعية/الفضائل/طرق الحساب غير مترجمة للغات الأخرى'),
    ('ف٢٣', 'دين', 'index.html ~2962', 'آية الكرسي بعد الصلاة: النسائي ١٣٣٦ — رقم مشكوك فيه. راجع "عمل اليوم والليلة"'),
    ('ف٢٤', 'دين', 'index.html ~2938', 'الصلاة على النبي ×10 صباحاً: "رواه الطبراني حسن صحيح" بدون رقم — يحتاج تحقق'),
    ('ف٢٥', 'بناء', 'build.gradle ~33', 'minifyEnabled false في release — فكّر في تفعيل ProGuard للطبقة الأصلية'),
    ('ف٢٦', 'بناء', 'لا يوجد', 'لا Content-Security-Policy — أضف <meta http-equiv="CSP"> لتقييد script-src'),
    ('ف٢٧', 'أداء', 'index.html ~4236', 'كود Ramadan (TIMES غير معرَّف) + PREMIUM_ENABLED=false = كود ميت — نظّفه'),
]

add_table(preferred_headers, preferred_rows, col_widths=[0.4, 0.8, 1.6, 3.5])
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — OK
# ════════════════════════════════════════════════════════════════════════════
heading('ثالثاً — سليم 🟢 (نتائج مطمئنة)', 1)

heading('القرآن الكريم', 2)
ok_quran = [
    'جميع 114 سورة: عدد الآيات في meta وفي النص متطابق تماماً',
    'سورة التوبة بدون بسملة ✓',
    'الفاتحة، الإخلاص، الفلق، الناس: نص وعدد صحيح ✓',
    'الرسم العثماني (ٱللَّهُ، همزة الوصل، السكون العثماني) مطبَّق باتساق ✓',
    'لا نقص ولا زيادة في أي سورة فُحصت ✓',
]
for item in ok_quran:
    bullet(item)

heading('الأذكار والأدعية والفضائل', 2)
ok_religion = [
    'جميع 64 ذكراً لها مصادر مع أرقام الأحاديث ✓',
    'جميع 39 دعاءً لها مصادر ✓',
    'جميع 44 فضيلة لها مصادر ✓',
    'دعاء ختم القرآن صُرِّح فيه صراحةً بـ "ليس حديثاً مرفوعاً" — شفافية ممتازة ✓',
    'لا أحاديث موضوعة أو مبالغ فيها في مكافآت العبادة ✓',
]
for item in ok_religion:
    bullet(item)

heading('حساب أوقات الصلاة', 2)
ok_prayer = [
    '5 طرق معتمدة بزوايا صحيحة: MWL (رابطة العالم الإسلامي)، ISNA، أم القرى، المصرية، كراتشي ✓',
    'حساب العصر: الجمهور (factor=1) والحنفي (factor=2) كلاهما صحيح ✓',
    'خوارزمية Meeus-based solar position — معيار مُعتمد ✓',
    'الطريقة الافتراضية: رابطة العالم الإسلامي — مناسب دولياً ✓',
]
for item in ok_prayer:
    bullet(item)

heading('الأمان', 2)
ok_security = [
    'قواعد Firebase RTDB: auth.uid === $uid — لا قراءة/كتابة عامة ✓',
    'allowMixedContent: false في capacitor.config.json ✓',
    'لا eval() في أي مكان بالملف ✓',
    'DOMParser يُستخدم لتحليل نتائج الحديث (لا innerHTML من الشبكة) ✓',
    'لا console.log ببيانات شخصية ✓',
    'جميع URLs الخارجية بـ HTTPS ✓',
]
for item in ok_security:
    bullet(item)

heading('الخصوصية', 2)
ok_privacy = [
    'الموقع الجغرافي: لا يغادر الجهاز (لا fetch للإحداثيات لأي خادم) ✓',
    'الكاميرا: معالجة محلية (لا رفع لأي إطار) ✓',
    'حذف الحساب: يحذف RTDB ثم Firebase Auth بترتيب صحيح مع معالجة requires-recent-login ✓',
    'إفصاحات الخصوصية في التطبيق + privacy.html + الكود متوافقة ✓',
]
for item in ok_privacy:
    bullet(item)

heading('البناء والامتثال', 2)
ok_build = [
    'targetSdk 35، compileSdk 35 ✓',
    'codemagic.yaml: prepare:www قبل cap sync في كلا الـ workflows ✓',
    '.nojekyll موجود لـ GitHub Pages ✓',
    'Firebase SDK محلي (./firebase/) في الـ script tags ✓',
    'PREMIUM_ENABLED=false — لا paywall يعيق الوصول ✓',
    'Apple Sign-In button: display:none افتراضياً، يظهر بعد كشف iOS ✓',
    'قسم تواصل ودعم: الـ 5 إجراءات مربوطة ✓',
    'جميع 6 NS*UsageDescription مكتملة في iOS Info.plist ✓',
    'جميع 191 مدخل في DICT لها ترجمات في اللغات الـ 5 ✓',
]
for item in ok_build:
    bullet(item)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — Regression table
# ════════════════════════════════════════════════════════════════════════════
heading('رابعاً — التحقق من الانحدارات (آخر 7 إصلاحات)', 1)

reg_headers = ['التغيير', 'الحالة', 'ملاحظة']
reg_rows = [
    ('Firebase SDK من ./firebase/ المحلي', '🟢 سليم', 'السطور 5650–5652 مؤكدة'),
    ('SW cache mujaddid-v43', '🟢 سليم', 'sw.js سطر 1، تنظيف الكاش القديم صحيح'),
    ('SW fallback للـ document فقط', '🔴 مكسور', 'خطأ أسبقية → المانع م١'),
    ('توجيه التعرف الصوتي عبر _isNative()', '🟢 سليم', 'مسار native plugin ومسار Web Speech كلاهما صحيح'),
    ('حذف الحساب fbDeleteAccount/_execDeleteAccount', '🟡 جزئي', 'التصميم صحيح لكن cred.idToken خطأ (م٣)'),
    ('DOMParser لنتائج الحديث', '🟢 سليم', 'searchHadith() لا يستخدم innerHTML من الشبكة'),
    ('Apple Sign-In مخفي على غير iOS', '🟢 سليم', 'display:none + كشف iOS صحيح'),
    ('قسم تواصل ودعم', '🟢 سليم', 'جميع الـ 5 إجراءات مربوطة'),
]
add_table(reg_headers, reg_rows, col_widths=[2.5, 1.2, 3.0])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — Priority fix plan
# ════════════════════════════════════════════════════════════════════════════
heading('خامساً — خطة الإصلاح بالأولوية', 1)

plan_headers = ['المرحلة', 'المانع', 'الأثر', 'الجهد']
plan_rows = [
    ('أولاً', 'م١ — SW operator precedence (سطر واحد)', 'التطبيق لا يعمل أوفلاين', 'دقائق'),
    ('أولاً', 'م٩ — Wake lock لا يُحرَّر', 'استنزاف البطارية', 'ساعة'),
    ('أولاً', 'م٣ — cred.idToken → cred.credential.idToken (كلمتان)', 'حذف الحساب مكسور', 'دقائق'),
    ('ثانياً', 'م١١ — حذف Facebook SDK', 'رفض Play Store', 'ساعة'),
    ('ثانياً', 'م١٢ — تطبيق Apple Sign-In الحقيقي', 'رفض App Store', 'يوم'),
    ('ثالثاً', 'م١٠ — استبدال JSONP بـ fetch()', 'ثغرة أمنية', 'يوم'),
    ('ثالثاً', 'م٢ — إزالة auth listener الداخلي', 'تعطّل تدريجي', 'ساعة'),
    ('رابعاً', 'م٤ — showDua() خارج شرط الإشعارات', 'دعاء لا يظهر', 'دقائق'),
    ('رابعاً', 'م٥ — buildQFlat() في requestIdleCallback', 'تجمّد واجهة', 'ساعة'),
    ('رابعاً', 'م٦ — TIMES → PT.data.timings.Maghrib', 'عداد رمضان صامت', 'دقائق'),
    ('رابعاً', 'م٧ — معالج زر الرجوع أندرويد', 'تجربة مستخدم مكسورة', 'يوم'),
    ('رابعاً', 'م٨ — e.message → textContent', 'XSS محتمل', 'دقائق'),
]
add_table(plan_headers, plan_rows, col_widths=[1.0, 2.8, 2.2, 1.0])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — Final verdict
# ════════════════════════════════════════════════════════════════════════════
heading('الحكم النهائي', 1)

p = doc.add_paragraph()
set_rtl(p)
run = p.add_run('🔴 التطبيق غير جاهز للرفع')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
set_run_rtl(run)

rtl_para(
    '12 مانعاً يجب إصلاحها. 8 منها إصلاحات سريعة (دقائق إلى ساعات). '
    'المانعان الأصعبان هما Apple Sign-In الحقيقي (م١٢) واستبدال JSONP (م١٠) '
    'ويحتاجان يوماً كاملاً لكل منهما.'
)
doc.add_paragraph()
rtl_para(
    'المحتوى الديني بجودة عالية: القرآن الكريم مضبوط، الأذكار موثَّقة، '
    'الفضائل والأدعية لها مصادر. نقطتان تحتاجان مراجعة بشرية متخصصة '
    '(آية الكرسي رقم النسائي، والصلاة على النبي ×10 مصدر الطبراني).'
)
doc.add_paragraph()

summary_headers = ['التصنيف', 'العدد']
summary_rows = [
    ('🔴 موانع للنشر', '12'),
    ('🟡 إصلاحات مُفضَّلة', '27'),
    ('🟢 سليم / مؤكد يعمل', '30+'),
    ('يحتاج مراجعة دينية بشرية', '2'),
]
add_table(summary_headers, summary_rows, col_widths=[3.5, 1.5])

doc.add_paragraph()
rtl_para('— نهاية التقرير —', bold=True)
rtl_para('DDC Business · المجدّد — طريق الأتقياء · 2026')

# ── Save ────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'audit-report.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
